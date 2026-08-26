"""Fill OATI letter DOCX template and append map/photo pages."""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run

MSK = ZoneInfo("Europe/Moscow")

TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "oati_letter.docx"
CAPTION_ICON_PATH = Path(__file__).resolve().parent / "templates" / "report.png"
MAP_WIDTH_CM_DOCX = 15.0
BODY_FONT_NAME = "Times New Roman"
BODY_FONT_SIZE = Pt(14)
MAP_PAGE_TITLE = "Ситуационный план места проведения земельных работ"
DEFAULT_DESCRIPTION = "Земляные работы при строительстве подземных коммуникаций."

# Private-use markers wrap auto-filled values so they render bold in Word.
_BOLD_OPEN = "\ue000"
_BOLD_CLOSE = "\ue001"

# Placeholders as they appear in the joined paragraph text (may be split across runs).
PH_DOC_DATE = "{cn{document.date(DateRU)(Concat:от | г. )}cn}"
PH_DOC_NUMBER = "{cn{document.number(Concat:№ |)}cn}"
PH_STREET = "{Улица на которой находится объект reports}"
PH_TODAY = "{Актуальное сегодняшнее число}"
PH_FID = "{fid из новой таблицы писем}"
PH_EXECUTOR = (
    "{Получаем из столбаца «Исполнитель», "
    "если работаем с задачей где этого нет можем ввести вручную или оставить пустым}"
)
PH_PHOTO_DT = (
    "{Дата и время берём из mggt_field.photos.created_at и нормализуем для РУ формата}"
)
PH_ADDRESS = (
    "{Получить адрес ближайшего здания к точке reports, "
    "использовать сторонние бесплатные сервисы, отображать только улицу и дом}"
)
PH_COORDS = "{координаты объекта reports}"
PH_ENG = (
    "{Получаем из столбаца data_mos.items62461_*.engineering_net_obj, "
    "если работаем с задачей где этого нет можем ввести вручную или оставить пустым}"
)
PH_DESCRIPTION = "{Ввод комментария вручную}"
PH_VIOLATION = "{Признаки незаконности из справочника}"
PH_PHOTO_COUNT = "{число выбранных фото}"
PH_SECTION_7 = "7. Признаки незаконности:"
SECTION_7_HEADER = "7. Данные Мосгоргеотрест:"
SECTION_8_HEADER = (
    "8. Данные, указывающие на признаки наличия административного правонарушения:"
)
STATION_UNDEFINED = "не определено"

# Soft line breaks lost during placeholder join — re-insert before these markers.
_LINE_BREAK_BEFORE = (
    "1. Сведения о производителе работ:",
    SECTION_7_HEADER,
    SECTION_8_HEADER,
)

# After join, ensure list content starts on the next line under the section header.
_LINE_BREAK_AFTER = (
    SECTION_7_HEADER,
    SECTION_8_HEADER,
)

# Subject line in letterhead table: restore soft break after join.
_SUBJECT_JOINED = "О предоставлении информацииоб инциденте"
_SUBJECT_SPLIT = "О предоставлении информации\nоб инциденте"


def format_ru_date(dt: datetime | None = None) -> str:
    value = dt or datetime.now(MSK)
    if value.tzinfo is None:
        value = value.replace(tzinfo=MSK)
    else:
        value = value.astimezone(MSK)
    return value.strftime("%d.%m.%Y")


def format_ru_datetime(value: str | datetime | None) -> str:
    """Legacy helper: date + time. Prefer ``format_ru_date_value`` for letter item 2."""
    if value is None or value == "":
        return ""
    if isinstance(value, datetime):
        dt = value
    else:
        text = str(value).strip()
        if not text:
            return ""
        try:
            dt = datetime.fromisoformat(text.replace("Z", "+00:00"))
        except ValueError:
            return text
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=MSK)
    else:
        dt = dt.astimezone(MSK)
    return dt.strftime("%d.%m.%Y %H:%M")


def format_ru_date_value(value: str | datetime | None) -> str:
    """Normalize ISO/datetime to ``ДД.ММ.ГГГГ`` (date only)."""
    if value is None or value == "":
        return ""
    if isinstance(value, datetime):
        dt = value
    else:
        text = str(value).strip()
        if not text:
            return ""
        try:
            dt = datetime.fromisoformat(text.replace("Z", "+00:00"))
        except ValueError:
            # Already a date-like string without time.
            if len(text) >= 10 and text[2] == "." and text[5] == ".":
                return text[:10]
            return text
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=MSK)
    else:
        dt = dt.astimezone(MSK)
    return dt.strftime("%d.%m.%Y")


def format_wgs84(lon: float, lat: float) -> str:
    return f"{lat:.6f}, {lon:.6f}"


def letter_download_filename(*, street: str, today: str, fid: int | str) -> str:
    """Download name: ``Об инциденте {street} от {today} №{fid}.docx``."""
    street_part = (street or "").strip() or "__________"
    for ch in '<>:"/\\|?*':
        street_part = street_part.replace(ch, "")
    street_part = " ".join(street_part.split())
    name = f"Об инциденте {street_part} от {today} №{fid}.docx"
    return " ".join(name.split())


def yandex_maps_url(lon: float, lat: float) -> str:
    return f"https://yandex.ru/maps/?pt={lon},{lat}&z=17&l=map"


def map_caption_text(scale: int, lat: float, lon: float) -> str:
    """Plain-text caption lines for the situational map (URL on its own line)."""
    coords = format_wgs84(lon, lat)
    url = yandex_maps_url(lon, lat)
    return (
        f"Масштаб 1:{scale}.\n"
        f"— место проведения земляных работ;\n"
        f"Координаты инцидента в WGS 84: {coords};\n"
        f"{url}"
    )


def photo_caption_label(index: int, *, banner: bool) -> str:
    kind = "Информационный щит" if banner else "Обзорное"
    return f"Фото {index} · {kind}"


def format_station_line(label: str, value: str | None) -> str:
    text = (value or "").strip() or STATION_UNDEFINED
    if text.lower().startswith(f"{label.lower()}:"):
        return text
    return f"{label}: {text}"


def format_mggt_block(sps: str | None, kgs: str | None) -> str:
    return f"{format_station_line('ТЗ ОПС', sps)}\n{format_station_line('КГС', kgs)}"


def format_violation_block(names: list[str]) -> str:
    """Join selected illegal-reason names as a bullet list (one item per line)."""
    cleaned = [n.strip() for n in names if (n or "").strip()]
    if not cleaned:
        return "__________"
    return "\n".join(f"• {name}" for name in cleaned)


def format_producer_block(customer: str, executor: str) -> str:
    """Section 1 body: Заказчик / Исполнитель lines; omit empty labels; bold values."""
    c = (customer or "").strip()
    e = (executor or "").strip()
    lines: list[str] = []
    if c:
        lines.append(f"Заказчик: {_mark_bold(c)}")
    if e:
        lines.append(f"Исполнитель: {_mark_bold(e)}")
    if not lines:
        return _mark_bold("__________")
    if len(lines) == 1:
        return lines[0]
    # Second line indented like the sample letter (tabs before «Исполнитель»).
    return f"{lines[0]}\n\t\t\t\t\t\t {lines[1]}"


def _mark_bold(value: str) -> str:
    """Wrap auto-filled value so ``_set_paragraph_text`` renders it bold."""
    return f"{_BOLD_OPEN}{value}{_BOLD_CLOSE}"


def _mark_bold_lines(value: str) -> str:
    """Bold each line separately (needed when value contains soft breaks)."""
    if "\n" not in value:
        return _mark_bold(value)
    return "\n".join(_mark_bold(line) if line else "" for line in value.split("\n"))


def _add_hyperlink(paragraph: Paragraph, url: str, text: str, *, font_size: Pt = Pt(9)) -> None:
    """Append a clickable hyperlink run to ``paragraph``."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), BODY_FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), BODY_FONT_NAME)
    r_fonts.set(qn("w:cs"), BODY_FONT_NAME)
    r_pr.append(r_fonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size.pt * 2)))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(font_size.pt * 2)))
    r_pr.append(sz_cs)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _ensure_structural_breaks(text: str) -> str:
    """Insert ``\\n`` before/after numbered section markers when missing."""
    result = text
    if _SUBJECT_JOINED in result:
        result = result.replace(_SUBJECT_JOINED, _SUBJECT_SPLIT, 1)
    for marker in _LINE_BREAK_BEFORE:
        idx = 0
        while True:
            pos = result.find(marker, idx)
            if pos < 0:
                break
            if pos > 0 and result[pos - 1] != "\n":
                result = result[:pos] + "\n" + result[pos:]
                pos += 1
            idx = pos + len(marker)
    for marker in _LINE_BREAK_AFTER:
        idx = 0
        while True:
            pos = result.find(marker, idx)
            if pos < 0:
                break
            end = pos + len(marker)
            if end < len(result) and result[end] != "\n":
                result = result[:end] + "\n" + result[end:]
                end += 1
            idx = end
    return result


def _apply_body_font(run: Run, *, bold: bool | None = None) -> None:
    """Force Times New Roman 14 on a run (incl. complex-script / East Asian)."""
    run.font.name = BODY_FONT_NAME
    run.font.size = BODY_FONT_SIZE
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), BODY_FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), BODY_FONT_NAME)
    r_fonts.set(qn("w:cs"), BODY_FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), BODY_FONT_NAME)
    # Mirror size for complex script (half-points: 14pt → 28).
    existing = r_pr.find(qn("w:szCs"))
    if existing is None:
        existing = OxmlElement("w:szCs")
        r_pr.append(existing)
    existing.set(qn("w:val"), "28")


def _clear_paragraph_runs(paragraph: Paragraph) -> None:
    p_el = paragraph._p
    for child in list(p_el):
        if child.tag == qn("w:r"):
            p_el.remove(child)


def _emit_marked_line(paragraph: Paragraph, line: str) -> Run | None:
    """Append runs for one line that may contain bold markers; return last run."""
    last: Run | None = None
    cursor = 0
    while cursor < len(line):
        open_at = line.find(_BOLD_OPEN, cursor)
        if open_at < 0:
            chunk = line[cursor:]
            if chunk:
                last = paragraph.add_run(chunk)
                _apply_body_font(last, bold=False)
            break
        if open_at > cursor:
            last = paragraph.add_run(line[cursor:open_at])
            _apply_body_font(last, bold=False)
        close_at = line.find(_BOLD_CLOSE, open_at + len(_BOLD_OPEN))
        if close_at < 0:
            # Unbalanced marker — emit remainder as normal text.
            last = paragraph.add_run(line[open_at + len(_BOLD_OPEN) :])
            _apply_body_font(last, bold=False)
            break
        bold_text = line[open_at + len(_BOLD_OPEN) : close_at]
        last = paragraph.add_run(bold_text)
        _apply_body_font(last, bold=True)
        cursor = close_at + len(_BOLD_CLOSE)
    return last


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph runs; ``\\n`` = soft break; ``\\ue000…\\ue001`` = bold value."""
    _clear_paragraph_runs(paragraph)
    parts = text.split("\n")
    for i, part in enumerate(parts):
        last = _emit_marked_line(paragraph, part)
        if i < len(parts) - 1:
            if last is None:
                last = paragraph.add_run("")
                _apply_body_font(last, bold=False)
            last.add_break(WD_BREAK.LINE)


def _normalize_body_paragraph_fonts(document: Document) -> None:
    """Ensure letter body paragraphs (not letterhead tables) are TNR 14; keep bold."""
    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            continue
        for run in paragraph.runs:
            was_bold = bool(run.bold)
            _apply_body_font(run, bold=was_bold if run.bold is not None else None)


def _replace_in_paragraph(paragraph: Paragraph, mapping: dict[str, str]) -> None:
    # Soft line breaks appear as "\n" in paragraph.text and split placeholders.
    full = paragraph.text.replace("\r", "").replace("\n", "")
    if not full:
        return
    new_text = full
    changed = False
    for key, value in mapping.items():
        if key in new_text:
            new_text = new_text.replace(key, value)
            changed = True
    if changed:
        _set_paragraph_text(paragraph, _ensure_structural_breaks(new_text))


def _iter_all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def fill_letter_template(
    *,
    street: str,
    today: str,
    fid: int | str,
    customer: str = "",
    executor: str = "",
    incident_datetime: str,
    address: str,
    coordinates: str,
    engineering: str,
    description: str,
    violation: str,
    photo_count: int = 0,
    sps: str = "",
    kgs: str = "",
) -> Document:
    if not TEMPLATE_PATH.is_file():
        raise FileNotFoundError(f"Letter template not found: {TEMPLATE_PATH}")

    document = Document(str(TEMPLATE_PATH))
    blank = "__________"
    desc = (description or "").strip() or DEFAULT_DESCRIPTION
    viol = (violation or "").strip() or blank
    producer = format_producer_block(customer, executor)
    sps_line = format_station_line("ТЗ ОПС", sps)
    kgs_line = format_station_line("КГС", kgs)
    section_7_8 = (
        f"{SECTION_7_HEADER}\n"
        f"{_mark_bold(sps_line)}\n"
        f"{_mark_bold(kgs_line)}\n"
        f"{SECTION_8_HEADER}"
    )
    # Bold only for auto-filled values inside sections 1–8.
    unique_map = {
        PH_DOC_DATE: f"от {today} г." if today else blank,
        PH_DOC_NUMBER: f"№ {fid}",
        PH_STREET: street or blank,
        PH_TODAY: today or blank,
        PH_FID: str(fid),
        PH_EXECUTOR: producer,
        PH_PHOTO_DT: _mark_bold(incident_datetime if incident_datetime else blank),
        PH_ADDRESS: _mark_bold(address if address else blank),
        PH_COORDS: _mark_bold(coordinates if coordinates else blank),
        PH_ENG: _mark_bold(engineering if engineering else blank),
        PH_DESCRIPTION: _mark_bold(desc),
        PH_SECTION_7: section_7_8,
        PH_VIOLATION: _mark_bold_lines(viol),
        PH_PHOTO_COUNT: str(int(photo_count)),
    }
    for paragraph in _iter_all_paragraphs(document):
        _replace_in_paragraph(paragraph, unique_map)
    _normalize_body_paragraph_fonts(document)
    return document


def _caption_icon_bytes() -> bytes | None:
    if not CAPTION_ICON_PATH.is_file():
        return None
    try:
        from PIL import Image

        icon = Image.open(CAPTION_ICON_PATH).convert("RGBA")
        buf = io.BytesIO()
        icon.save(buf, format="PNG")
        return buf.getvalue()
    except OSError:
        return None


def _add_caption_legend_icon(paragraph: Paragraph) -> None:
    raw = _caption_icon_bytes()
    run = paragraph.add_run()
    if raw:
        try:
            run.add_picture(io.BytesIO(raw), height=Cm(0.5))
            return
        except Exception:
            pass
    run.text = "●"
    run.font.name = BODY_FONT_NAME
    run.font.size = Pt(9)


def _apply_caption_font(run: Run) -> None:
    run.font.name = BODY_FONT_NAME
    run.font.size = Pt(9)


def _image_is_portrait(image_bytes: bytes) -> bool:
    try:
        from PIL import Image

        with Image.open(io.BytesIO(image_bytes)) as img:
            width, height = img.size
        return height > width
    except OSError:
        return False


def _pack_photo_pages(
    photos: list[tuple[bytes, str]],
) -> list[list[tuple[bytes, str]]]:
    """Greedy pack: 3 landscape per page, 2 if the page has a portrait."""
    pages: list[list[tuple[bytes, str]]] = []
    current: list[tuple[bytes, str, bool]] = []

    def flush() -> None:
        if current:
            pages.append([(blob, label) for blob, label, _orient in current])
            current.clear()

    for image_bytes, label in photos:
        portrait = _image_is_portrait(image_bytes)
        has_portrait = any(item[2] for item in current)
        limit = 2 if (portrait or has_portrait) else 3
        if current and len(current) >= limit:
            flush()
        current.append((image_bytes, label, portrait))
    flush()
    return pages


def _add_fitted_photo(run: Run, image_bytes: bytes, *, max_height_cm: float) -> None:
    from PIL import Image

    with Image.open(io.BytesIO(image_bytes)) as img:
        width, height = img.size
    if width <= 0 or height <= 0:
        raise ValueError("invalid image size")
    max_width_cm = 14.0
    height_cm = max_height_cm
    width_cm = height_cm * (width / height)
    if width_cm > max_width_cm:
        width_cm = max_width_cm
        height_cm = width_cm * (height / width)
    run.add_picture(io.BytesIO(image_bytes), width=Cm(width_cm), height=Cm(height_cm))


def _add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def append_map_page(
    document: Document,
    map_png: bytes,
    title: str = MAP_PAGE_TITLE,
    *,
    scale: int = 1000,
    lon: float | None = None,
    lat: float | None = None,
) -> None:
    _add_page_break(document)
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    _apply_body_font(run)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(map_png), width=Cm(MAP_WIDTH_CM_DOCX))

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if lon is not None and lat is not None:
        url = yandex_maps_url(lon, lat)
        coords = format_wgs84(lon, lat)
        r1 = caption.add_run(f"Масштаб 1:{scale}.")
        _apply_caption_font(r1)
        r1.add_break(WD_BREAK.LINE)
        _add_caption_legend_icon(caption)
        r2 = caption.add_run(" — место проведения земляных работ;")
        _apply_caption_font(r2)
        r2.add_break(WD_BREAK.LINE)
        r3 = caption.add_run(f"Координаты инцидента в WGS 84: {coords};")
        _apply_caption_font(r3)
        r3.add_break(WD_BREAK.LINE)
        _add_hyperlink(caption, url, url, font_size=Pt(9))
    else:
        r1 = caption.add_run(f"Масштаб 1:{scale}.")
        _apply_caption_font(r1)
        r1.add_break(WD_BREAK.LINE)
        _add_caption_legend_icon(caption)
        r2 = caption.add_run(" — место проведения земляных работ.")
        _apply_caption_font(r2)


def append_photo_pages(
    document: Document,
    photos: list[tuple[bytes, str]],
) -> None:
    """Append photos starting on a new page; 3 landscape or 2 if a portrait is present."""
    if not photos:
        _add_page_break(document)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Фотофиксация: фотографии не выбраны.")
        _apply_body_font(run)
        return

    pages = _pack_photo_pages(photos)
    for page_index, page_photos in enumerate(pages):
        _add_page_break(document)
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title = "Фотофиксация" if page_index == 0 else "Фотофиксация (продолжение)"
        run = heading.add_run(title)
        run.bold = True
        _apply_body_font(run)

        has_portrait = any(_image_is_portrait(blob) for blob, _label in page_photos)
        max_height_cm = 10.0 if (has_portrait or len(page_photos) <= 2) else 6.7

        for image_bytes, label in page_photos:
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
            label_run = caption.add_run(label)
            _apply_body_font(label_run)

            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_run = paragraph.add_run()
            try:
                _add_fitted_photo(pic_run, image_bytes, max_height_cm=max_height_cm)
            except Exception:
                paragraph.add_run(" [не удалось вставить изображение] ")


def document_to_bytes(document: Document) -> bytes:
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
