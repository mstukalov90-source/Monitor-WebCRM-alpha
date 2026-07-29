"""Tests for OATI letter helpers: placeholders, map scale, geocode, lookups."""

from __future__ import annotations

import io
import unittest
from unittest.mock import MagicMock, patch
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from app.auth.deps import require_manager_or_admin
from app.auth.session import UserSession
from app.letters.docx_fill import (
    BODY_FONT_NAME,
    BODY_FONT_SIZE,
    DEFAULT_DESCRIPTION,
    PH_DESCRIPTION,
    PH_DOC_DATE,
    PH_STREET,
    PH_VIOLATION,
    TEMPLATE_PATH,
    append_map_page,
    append_photo_pages,
    document_to_bytes,
    fill_letter_template,
    format_producer_block,
    format_ru_date,
    format_ru_date_value,
    format_ru_datetime,
    format_violation_block,
    format_wgs84,
    letter_download_filename,
    map_caption_text,
    photo_caption_label,
    yandex_maps_url,
)
from app.letters.geocode import (
    HOUSE_SEARCH_RADIUS_M,
    GeocodeResult,
    format_street_house,
    reverse_geocode_parts,
)
from app.letters.map_image import (
    ALLOWED_MAP_SCALES,
    DEFAULT_MAP_SCALE,
    GROUND_WIDTH_M,
    MAP_SCALE,
    classify_geometry_visibility,
    ground_width_m,
    map_bbox_mercator,
    normalize_map_scale,
)
from app.letters.oati import (
    LetterError,
    _lookup_customer,
    _lookup_engineering,
    _lookup_executor,
    _lookup_mos_simple_address,
    _validate_photo_ids,
    _validate_violation_names,
    merge_engineering_values,
    pick_default_address,
    resolve_incident_datetime,
)
from app.crm.schemas import OatiLetterGenerateRequest
from pydantic import ValidationError
from fastapi import HTTPException


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _docx_text(data: bytes) -> str:
    root = ET.fromstring(ZipFile(io.BytesIO(data)).read("word/document.xml"))
    return "".join((t.text or "") for t in root.findall(".//w:t", NS))


def _docx_image_count(data: bytes) -> int:
    root = ET.fromstring(ZipFile(io.BytesIO(data)).read("word/document.xml"))
    drawings = root.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
    )
    return len(drawings)


class DocxFillTests(unittest.TestCase):
    def test_replaces_new_template_placeholders(self) -> None:
        doc = fill_letter_template(
            street="ул. Ленина",
            today="23.07.2026",
            fid=7,
            customer="АО Заказчик",
            executor="ООО Строй",
            incident_datetime="22.07.2026",
            address="ул. Ленина, 10",
            coordinates="55.800000, 37.500000",
            engineering="не определено",
            description="описание А",
            violation=format_violation_block(
                ["Отсутствие КГС", "Отсутствие уведомления"]
            ),
            photo_count=3,
        )
        from PIL import Image

        buf = io.BytesIO()
        Image.new("RGB", (40, 40), (180, 180, 180)).save(buf, format="PNG")
        png = buf.getvalue()
        append_map_page(doc, png, scale=2000, lon=37.5, lat=55.8)
        append_photo_pages(doc, [(png, "Фото 1 · Обзорное фото")])
        data = document_to_bytes(doc)
        text = _docx_text(data)

        self.assertNotIn(PH_DOC_DATE, text)
        self.assertNotIn(PH_STREET, text)
        self.assertNotIn(PH_DESCRIPTION, text)
        self.assertNotIn(PH_VIOLATION, text)
        self.assertIn("от 23.07.2026 г.", text)
        self.assertIn("№ 7", text)
        self.assertIn("ул. Ленина", text)
        self.assertIn("Заказчик:", text)
        self.assertIn("АО Заказчик", text)
        self.assertIn("Исполнитель:", text)
        self.assertIn("ООО Строй", text)
        self.assertIn("описание А", text)
        self.assertIn("незаконности земляных работ при строительстве инженерных коммуникаций", text)
        self.assertIn("Описание характера работ", text)
        self.assertIn("Признаки незаконности", text)
        self.assertIn("• Отсутствие КГС", text)
        self.assertIn("• Отсутствие уведомления", text)
        self.assertIn("ОАТИ города Москвы", text)
        self.assertIn("Фотофиксация: 3 экз.", text)
        self.assertIn("Ситуационный план места проведения земельных работ", text)
        self.assertIn("О предоставлении информации", text)
        self.assertIn("об инциденте ул. Ленина от 23.07.2026 №7", text)
        self.assertGreaterEqual(_docx_image_count(data), 2)

        # Soft breaks before items 1 and 7 must survive fill; list under п.7.
        joined = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("\n1. Сведения о производителе работ:", joined)
        self.assertIn("Заказчик: АО Заказчик", joined)
        self.assertIn("Исполнитель: ООО Строй", joined)
        self.assertIn("\n7. Признаки незаконности:\n• Отсутствие КГС", joined)

        # Auto-filled values in п.1–7 are semi-bold; labels and outside text stay regular.
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                chunk = run.text or ""
                if "описание А" in chunk or chunk.strip().startswith("• "):
                    self.assertTrue(run.bold, msg=repr(chunk[:60]))
                if "АО Заказчик" in chunk or "ООО Строй" in chunk:
                    self.assertTrue(run.bold, msg=repr(chunk[:60]))
                if chunk.startswith("6. Описание") or chunk.startswith("7. Признаки"):
                    self.assertFalse(run.bold, msg=repr(chunk[:60]))
                if "Заказчик:" in chunk and "АО" not in chunk:
                    self.assertFalse(run.bold, msg=repr(chunk[:60]))

        # Title / subject / appendix auto-fill must not be bold.
        for paragraph in doc.paragraphs:
            if not paragraph.text.startswith("Отчёт об инциденте"):
                continue
            for run in paragraph.runs:
                chunk = (run.text or "").strip()
                if not chunk:
                    continue
                # Stop at section 1 — values after that may be bold (п.1).
                if "1. Сведения" in chunk or chunk.startswith("Заказчик") or chunk.startswith("Исполнитель"):
                    break
                if chunk in ("ул. Ленина", "23.07.2026", "7", "от", "№", ":"):
                    self.assertFalse(bool(run.bold), msg=repr(chunk))

        for paragraph in doc.paragraphs:
            if "Фотофиксация:" not in paragraph.text:
                continue
            for run in paragraph.runs:
                if "3" in (run.text or ""):
                    self.assertFalse(bool(run.bold), msg=repr(run.text))

        # Subject soft break in letterhead table.
        subject_joined = "\n".join(
            p.text for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs
        )
        self.assertIn("О предоставлении информации\nоб инциденте ул. Ленина", subject_joined)

        # Body text (report + appendix) must be Times New Roman 14.
        body_markers = (
            "ГБУ «Мосгоргеотрест»",
            "Отчёт об инциденте",
            "1. Сведения о производителе работ",
            "2. Дата и время",
            "7. Признаки незаконности",
            "Приложение:",
            "1. Ситуационный план:",
        )
        for paragraph in doc.paragraphs:
            if not any(m in paragraph.text for m in body_markers):
                continue
            for run in paragraph.runs:
                if not (run.text or "").strip():
                    continue
                self.assertEqual(run.font.name, BODY_FONT_NAME, msg=repr(run.text[:40]))
                self.assertEqual(
                    run.font.size.pt if run.font.size else None,
                    BODY_FONT_SIZE.pt,
                    msg=repr(run.text[:40]),
                )

    def test_map_caption_and_yandex_link(self) -> None:
        self.assertEqual(
            yandex_maps_url(37.5, 55.8),
            "https://yandex.ru/maps/?pt=37.5,55.8&z=17&l=map",
        )
        caption = map_caption_text(1000, 55.8, 37.5)
        self.assertIn("Масштаб 1:1000", caption)
        self.assertIn("Красный знак — место проведения земляных работ", caption)
        self.assertNotIn("разрытие", caption)
        self.assertIn("55.800000, 37.500000", caption)
        self.assertIn("https://yandex.ru/maps/?pt=37.5,55.8&z=17&l=map", caption)

        from PIL import Image

        doc = fill_letter_template(
            street="ул. А",
            today="24.07.2026",
            fid=1,
            executor="",
            incident_datetime="",
            address="",
            coordinates="",
            engineering="",
            description="",
            violation="",
            photo_count=0,
        )
        buf = io.BytesIO()
        Image.new("RGB", (20, 20), (100, 100, 100)).save(buf, format="PNG")
        append_map_page(doc, buf.getvalue(), scale=5000, lon=37.5, lat=55.8)
        data = document_to_bytes(doc)
        text = _docx_text(data)
        self.assertIn("Масштаб 1:5000", text)
        self.assertIn("Красный знак — место проведения земляных работ", text)
        self.assertIn("yandex.ru/maps", text)
        self.assertIn(DEFAULT_DESCRIPTION, text)

        # Hyperlink relationship must exist in the package.
        with ZipFile(io.BytesIO(data)) as zf:
            rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        self.assertIn("yandex.ru/maps", rels)

    def test_format_helpers(self) -> None:
        self.assertEqual(format_wgs84(37.5, 55.8), "55.800000, 37.500000")
        self.assertRegex(format_ru_date(), r"\d{2}\.\d{2}\.\d{4}")
        self.assertEqual(format_ru_datetime("2026-07-22T12:30:00+03:00"), "22.07.2026 12:30")
        self.assertEqual(format_ru_date_value("2026-07-22T12:30:00+03:00"), "22.07.2026")
        self.assertEqual(
            letter_download_filename(street="улица Фомичёвой", today="27.07.2026", fid=15),
            "Об инциденте улица Фомичёвой от 27.07.2026 №15.docx",
        )
        self.assertEqual(
            letter_download_filename(street='ул. A/B:C', today="01.01.2026", fid=1),
            "Об инциденте ул. ABC от 01.01.2026 №1.docx",
        )
        self.assertEqual(photo_caption_label(1, banner=True), "Фото 1 · Информационный щит")
        self.assertEqual(photo_caption_label(2, banner=False), "Фото 2 · Обзорное фото")
        self.assertEqual(
            format_violation_block(["A", "B"]),
            "• A\n• B",
        )
        self.assertEqual(format_violation_block([]), "__________")
        self.assertEqual(
            DEFAULT_DESCRIPTION,
            "Земляные работы при строительстве подземных коммуникаций.",
        )
        both = format_producer_block("АО Заказчик", "ООО Исполнитель")
        self.assertIn("Заказчик:", both)
        self.assertIn("АО Заказчик", both)
        self.assertIn("Исполнитель:", both)
        self.assertIn("ООО Исполнитель", both)
        self.assertIn("\n", both)
        only_c = format_producer_block("АО Заказчик", "")
        self.assertIn("Заказчик:", only_c)
        self.assertNotIn("Исполнитель:", only_c)
        only_e = format_producer_block("", "ООО Исполнитель")
        self.assertIn("Исполнитель:", only_e)
        self.assertNotIn("Заказчик:", only_e)
        empty = format_producer_block("", "")
        self.assertIn("__________", empty)
        self.assertNotIn("Заказчик:", empty)
        self.assertNotIn("Исполнитель:", empty)

    def test_template_has_no_empty_tables(self) -> None:
        from docx import Document

        doc = Document(str(TEMPLATE_PATH))
        self.assertEqual(len(doc.tables), 2)
        for table in doc.tables:
            text = "\n".join(
                p.text for row in table.rows for cell in row.cells for p in cell.paragraphs
            ).strip()
            self.assertTrue(text, msg="empty table left in letter template")
        self.assertIn("ОАТИ города Москвы", doc.tables[0].cell(0, 1).text)


class MapScaleTests(unittest.TestCase):
    def test_ground_extent_matches_scale_1000(self) -> None:
        self.assertEqual(MAP_SCALE, 1000)
        self.assertEqual(DEFAULT_MAP_SCALE, 1000)
        self.assertEqual(GROUND_WIDTH_M, 160.0)
        self.assertEqual(ground_width_m(1000), 160.0)
        self.assertEqual(ground_width_m(2000), 320.0)
        self.assertEqual(ground_width_m(5000), 800.0)
        self.assertEqual(ground_width_m(10000), 1600.0)
        self.assertEqual(ALLOWED_MAP_SCALES, (1000, 2000, 5000, 10000))

    def test_report_marker_asset_loads(self) -> None:
        from app.letters.map_image import REPORT_MARKER_PATH, REPORT_MARKER_WIDTH_PX, _load_report_marker

        self.assertTrue(REPORT_MARKER_PATH.is_file())
        icon = _load_report_marker()
        self.assertEqual(icon.mode, "RGBA")
        self.assertEqual(icon.size[0], REPORT_MARKER_WIDTH_PX)

    def test_bbox_width_scales(self) -> None:
        lon, lat = 37.5, 55.8
        for scale, expected in ((1000, 160.0), (2000, 320.0), (5000, 800.0), (10000, 1600.0)):
            minx, miny, maxx, maxy = map_bbox_mercator(lon, lat, scale=scale)
            self.assertAlmostEqual(maxx - minx, expected, places=3, msg=f"scale={scale}")
            self.assertAlmostEqual(maxy - miny, expected, places=3, msg=f"scale={scale}")

    def test_bbox_centered_on_report(self) -> None:
        lon, lat = 37.5, 55.8
        minx, miny, maxx, maxy = map_bbox_mercator(lon, lat)
        self.assertAlmostEqual(maxx - minx, GROUND_WIDTH_M, places=3)
        self.assertAlmostEqual(maxy - miny, GROUND_WIDTH_M, places=3)

    def test_normalize_map_scale_rejects_invalid(self) -> None:
        self.assertEqual(normalize_map_scale(None), 1000)
        self.assertEqual(normalize_map_scale(2000), 2000)
        with self.assertRaises(ValueError):
            normalize_map_scale(1500)

    def test_generate_request_rejects_invalid_scale(self) -> None:
        with self.assertRaises(ValidationError):
            OatiLetterGenerateRequest(map_scale=1234)
        req = OatiLetterGenerateRequest(map_scale=5000)
        self.assertEqual(req.map_scale, 5000)

    def test_geometry_visibility_clipping_states(self) -> None:
        center_lon, center_lat = 37.5, 55.8
        inside = {"type": "Point", "coordinates": [center_lon, center_lat]}
        outside = {"type": "Point", "coordinates": [center_lon + 0.05, center_lat + 0.05]}
        partial = {
            "type": "LineString",
            "coordinates": [
                [center_lon, center_lat],
                [center_lon + 0.05, center_lat + 0.05],
            ],
        }
        self.assertEqual(classify_geometry_visibility(inside, center_lon, center_lat), "inside")
        self.assertEqual(classify_geometry_visibility(outside, center_lon, center_lat), "outside")
        self.assertEqual(classify_geometry_visibility(partial, center_lon, center_lat), "partial")
        self.assertEqual(classify_geometry_visibility(None, center_lon, center_lat), "missing")


class GeocodeFormatTests(unittest.TestCase):
    def test_street_house_from_nominatim(self) -> None:
        self.assertEqual(
            format_street_house({"road": "ул. Ленина", "house_number": "10"}),
            "ул. Ленина, 10",
        )
        self.assertIsNone(format_street_house({}))
        self.assertEqual(format_street_house({"road": "Тверская"}), "Тверская")

    def test_reverse_parts_prefers_house_number(self) -> None:
        from app.config import Settings

        settings = Settings(nominatim_url="https://example.invalid/reverse")
        reverse_payload = {
            "address": {"road": "ул. Ленина", "house_number": "12к1"},
        }
        with patch("app.letters.geocode._nominatim_get", return_value=reverse_payload):
            result = reverse_geocode_parts(37.5, 55.8, settings)
        self.assertEqual(result.street, "ул. Ленина")
        self.assertEqual(result.address, "ул. Ленина, 12к1")
        self.assertTrue(result.has_house)

    def test_reverse_parts_searches_nearby_when_house_missing(self) -> None:
        import math
        from app.config import Settings
        from urllib.parse import parse_qs, urlparse

        settings = Settings(nominatim_url="https://example.invalid/reverse")
        search_urls: list[str] = []

        def fake_get(url: str, _settings: Settings):
            if "/search" in url:
                search_urls.append(url)
                return [
                    {
                        "lat": "55.80001",
                        "lon": "37.50001",
                        "address": {"road": "ул. Ленина", "house_number": "5"},
                    }
                ]
            return {"address": {"road": "ул. Ленина"}}

        with patch("app.letters.geocode._nominatim_get", side_effect=fake_get):
            result = reverse_geocode_parts(37.5, 55.8, settings)
        self.assertEqual(result.address, "ул. Ленина, 5")
        self.assertTrue(result.has_house)
        self.assertEqual(HOUSE_SEARCH_RADIUS_M, 250.0)
        self.assertTrue(search_urls)
        qs = parse_qs(urlparse(search_urls[0]).query)
        self.assertEqual(qs.get("limit"), ["20"])
        left, top, right, bottom = (float(x) for x in qs["viewbox"][0].split(","))
        half_lat = (top - bottom) / 2
        expected_lat = HOUSE_SEARCH_RADIUS_M / 111_320.0
        self.assertAlmostEqual(half_lat, expected_lat, places=6)
        expected_lon = HOUSE_SEARCH_RADIUS_M / (111_320.0 * max(0.2, math.cos(math.radians(55.8))))
        self.assertAlmostEqual((right - left) / 2, expected_lon, places=6)

    def test_reverse_geocode_returns_empty_on_network_error(self) -> None:
        from app.config import Settings

        settings = Settings(nominatim_url="https://example.invalid/reverse")
        with patch("app.letters.geocode.urllib.request.urlopen", side_effect=OSError("down")):
            result = reverse_geocode_parts(37.5, 55.8, settings)
        self.assertEqual(result, GeocodeResult())


class SourceLookupTests(unittest.TestCase):
    def test_executor_from_source_general_contractor(self) -> None:
        conn = MagicMock()
        record = MagicMock(key="task-1")
        with patch(
            "app.letters.oati._lookup_source_feature",
            return_value={"attributes": {"general_contractor": "ООО Ромашка"}},
        ):
            self.assertEqual(_lookup_executor(conn, record, {}), "ООО Ромашка")

    def test_executor_empty_when_source_missing(self) -> None:
        conn = MagicMock()
        record = MagicMock(key="task-1")
        with patch("app.letters.oati._lookup_source_feature", return_value={"attributes": {}}):
            self.assertEqual(_lookup_executor(conn, record, {}), "")

    def test_customer_from_source_fields(self) -> None:
        conn = MagicMock()
        record = MagicMock(key="task-1")
        with patch(
            "app.letters.oati._lookup_source_feature",
            return_value={"attributes": {"customer_construction": "ПАО МОЭК"}},
        ):
            self.assertEqual(_lookup_customer(conn, record, {}), "ПАО МОЭК")
        with patch(
            "app.letters.oati._lookup_source_feature",
            return_value={"attributes": {"balanceholder": "АО Баланс"}},
        ):
            self.assertEqual(_lookup_customer(conn, record, {}), "АО Баланс")
        with patch(
            "app.letters.oati._lookup_source_feature",
            return_value={"attributes": {"customer": "ООО Клиент"}},
        ):
            self.assertEqual(_lookup_customer(conn, record, {}), "ООО Клиент")

    def test_fill_omits_empty_producer_labels(self) -> None:
        doc = fill_letter_template(
            street="ул. А",
            today="24.07.2026",
            fid=1,
            customer="Только заказчик",
            executor="",
            incident_datetime="",
            address="",
            coordinates="",
            engineering="",
            description="",
            violation="",
            photo_count=0,
        )
        joined = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Заказчик: Только заказчик", joined)
        self.assertNotIn("Исполнитель:", joined)

        doc2 = fill_letter_template(
            street="ул. А",
            today="24.07.2026",
            fid=1,
            customer="",
            executor="Только исполнитель",
            incident_datetime="",
            address="",
            coordinates="",
            engineering="",
            description="",
            violation="",
            photo_count=0,
        )
        joined2 = "\n".join(p.text for p in doc2.paragraphs)
        self.assertIn("Исполнитель: Только исполнитель", joined2)
        self.assertNotIn("Заказчик:", joined2)

    def test_engineering_from_engineering_net_obj_not_type(self) -> None:
        conn = MagicMock()
        record = MagicMock(key="task-1", type="АВР")
        with patch(
            "app.letters.oati._lookup_source_feature",
            return_value={"attributes": {"engineering_net_obj": "теплосеть", "type": "ignore"}},
        ):
            self.assertEqual(_lookup_engineering(conn, record, {}), "теплосеть")

    def test_engineering_empty_when_only_task_type_present(self) -> None:
        conn = MagicMock()
        record = MagicMock(key="task-1", type="Разрытия")
        with patch(
            "app.letters.oati._lookup_source_feature",
            return_value={"attributes": {"something_else": "x"}},
        ):
            self.assertEqual(_lookup_engineering(conn, record, {}), "")

    def test_engineering_merges_report_comm_type(self) -> None:
        conn = MagicMock()
        record = MagicMock(key="task-1")
        with patch(
            "app.letters.oati._lookup_source_feature",
            return_value={"attributes": {"engineering_net_obj": "Тепловая сеть"}},
        ):
            self.assertEqual(
                _lookup_engineering(conn, record, {}, report_comm_type="Теплосеть"),
                "Тепловая сеть, Теплосеть",
            )

    def test_engineering_uses_comm_type_when_source_empty(self) -> None:
        conn = MagicMock()
        record = MagicMock(key="task-1")
        with patch(
            "app.letters.oati._lookup_source_feature",
            return_value={"attributes": {}},
        ):
            self.assertEqual(
                _lookup_engineering(conn, record, {}, report_comm_type="Газопровод"),
                "Газопровод",
            )

    def test_merge_engineering_values(self) -> None:
        self.assertEqual(merge_engineering_values("", ""), "")
        self.assertEqual(merge_engineering_values("A", ""), "A")
        self.assertEqual(merge_engineering_values("", "B"), "B")
        self.assertEqual(merge_engineering_values("Теплосеть", "теплосеть"), "Теплосеть")
        self.assertEqual(merge_engineering_values("A", "B"), "A, B")


class IncidentDatetimeTests(unittest.TestCase):
    def test_prefers_taken_at_over_created_at(self) -> None:
        from app.photos.field_photo import FieldPhotoItem

        photos = [
            FieldPhotoItem(
                id=1,
                file_path="a.jpg",
                banner=False,
                created_at="2026-07-23T16:40:00+03:00",
                photo_key=None,
                username=None,
                taken_at=None,
            ),
            FieldPhotoItem(
                id=2,
                file_path="b.jpg",
                banner=False,
                created_at="2026-07-23T10:00:00+03:00",
                photo_key=None,
                username=None,
                taken_at="2026-07-22T12:30:00+03:00",
            ),
        ]
        self.assertEqual(resolve_incident_datetime(photos), "22.07.2026")

    def test_falls_back_to_created_at_when_taken_at_null(self) -> None:
        from app.photos.field_photo import FieldPhotoItem

        photos = [
            FieldPhotoItem(
                id=1,
                file_path="a.jpg",
                banner=False,
                created_at="2026-07-23T16:40:00+03:00",
                photo_key=None,
                username=None,
                taken_at=None,
            ),
        ]
        self.assertEqual(resolve_incident_datetime(photos), "23.07.2026")

    def test_preferred_ids_order_for_taken_at(self) -> None:
        from app.photos.field_photo import FieldPhotoItem

        photos = [
            FieldPhotoItem(
                id=1,
                file_path="a.jpg",
                banner=False,
                created_at=None,
                photo_key=None,
                username=None,
                taken_at="2026-07-21T09:00:00+03:00",
            ),
            FieldPhotoItem(
                id=2,
                file_path="b.jpg",
                banner=False,
                created_at=None,
                photo_key=None,
                username=None,
                taken_at="2026-07-22T11:00:00+03:00",
            ),
        ]
        self.assertEqual(
            resolve_incident_datetime(photos, preferred_ids=[2, 1]),
            "22.07.2026",
        )


class ViolationNamesTests(unittest.TestCase):
    def test_accepts_known_names(self) -> None:
        conn = MagicMock()
        with patch(
            "app.letters.oati._fetch_illegal_reason_names",
            return_value=["A", "B", "C"],
        ):
            self.assertEqual(_validate_violation_names(conn, ["B", "A", "B"]), ["B", "A"])

    def test_rejects_unknown_names(self) -> None:
        conn = MagicMock()
        with patch(
            "app.letters.oati._fetch_illegal_reason_names",
            return_value=["A"],
        ):
            with self.assertRaises(LetterError) as ctx:
                _validate_violation_names(conn, ["A", "X"])
            self.assertEqual(ctx.exception.status_code, 422)
            self.assertIn("X", str(ctx.exception))

    def test_engineering_special_values_passthrough(self) -> None:
        # Form may send these radio values as engineering string.
        self.assertEqual(merge_engineering_values("не определено", ""), "не определено")
        self.assertEqual(merge_engineering_values("отсутствует", ""), "отсутствует")


class MosAddressLookupTests(unittest.TestCase):
    def test_contains_hit(self) -> None:
        conn = MagicMock()
        cur = MagicMock()
        cur.fetchone.return_value = ("Енисейская улица, дом 39А",)
        conn.cursor.return_value.__enter__.return_value = cur
        self.assertEqual(_lookup_mos_simple_address(conn, 37.5, 55.8), "Енисейская улица, дом 39А")
        self.assertEqual(cur.execute.call_count, 1)

    def test_contains_miss_falls_back_to_nearest(self) -> None:
        conn = MagicMock()
        cur = MagicMock()
        cur.fetchone.side_effect = [None, ("Соловьиный проезд, дом 2",)]
        conn.cursor.return_value.__enter__.return_value = cur
        self.assertEqual(_lookup_mos_simple_address(conn, 37.5, 55.8), "Соловьиный проезд, дом 2")
        self.assertEqual(cur.execute.call_count, 2)
        nearest_sql = cur.execute.call_args_list[1][0][0]
        self.assertIn("<->", nearest_sql)

    def test_db_error_returns_empty(self) -> None:
        conn = MagicMock()
        conn.cursor.side_effect = OSError("db down")
        self.assertEqual(_lookup_mos_simple_address(conn, 37.5, 55.8), "")

    def test_pick_default_address_prefers_geocode_with_house(self) -> None:
        self.assertEqual(
            pick_default_address(
                address_geocode="ул. А, 1",
                address_mos="ул. Б, дом 2",
                address_has_house=True,
            ),
            "ул. А, 1",
        )

    def test_pick_default_address_prefers_mos_without_house(self) -> None:
        self.assertEqual(
            pick_default_address(
                address_geocode="ул. А",
                address_mos="ул. Б, дом 2",
                address_has_house=False,
            ),
            "ул. Б, дом 2",
        )

    def test_pick_default_address_geocode_when_mos_empty(self) -> None:
        self.assertEqual(
            pick_default_address(
                address_geocode="ул. А",
                address_mos="",
                address_has_house=False,
            ),
            "ул. А",
        )


class PhotoValidationTests(unittest.TestCase):
    def test_rejects_foreign_photo_ids(self) -> None:
        conn = MagicMock()
        photo = MagicMock()
        photo.id = 10
        with patch(
            "app.letters.oati.fetch_field_photos",
            return_value=MagicMock(photos=[photo]),
        ):
            with self.assertRaises(LetterError) as ctx:
                _validate_photo_ids(conn, "task", 1, [10, 99])
            self.assertIn("99", str(ctx.exception))

    def test_dedupes_preserving_order(self) -> None:
        conn = MagicMock()
        p1, p2 = MagicMock(id=1), MagicMock(id=2)
        with patch(
            "app.letters.oati.fetch_field_photos",
            return_value=MagicMock(photos=[p1, p2]),
        ):
            self.assertEqual(_validate_photo_ids(conn, "task", 1, [2, 1, 2]), [2, 1])


class RbacTests(unittest.TestCase):
    def test_require_manager_or_admin_rejects_office(self) -> None:
        user = UserSession(uuid="u", login="office1", role="office", work_zones=[])
        with self.assertRaises(HTTPException) as ctx:
            require_manager_or_admin(user)
        self.assertEqual(ctx.exception.status_code, 403)

    def test_require_manager_or_admin_allows_manager(self) -> None:
        user = UserSession(uuid="u", login="mgr", role="manager", work_zones=[1])
        self.assertEqual(require_manager_or_admin(user), user)


if __name__ == "__main__":
    unittest.main()
